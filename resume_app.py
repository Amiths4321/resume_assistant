# resume_app.py
# streamlit run resume_app.py

import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import streamlit as st
from pathlib import Path
from resume_parser import extract_pdf_text
from resume_ai     import (
    score_job_fit, rewrite_resume,
    generate_cover_letter, predict_interview_questions,
    improve_resume
)

st.set_page_config(
    page_title="AI Resume Assistant",
    page_icon="📋",
    layout="wide"
)

Path("outputs").mkdir(exist_ok=True)

# ── Session state ─────────────────────────────────────────────────────────────
if "resume_text" not in st.session_state: st.session_state.resume_text = ""
if "jd_text"     not in st.session_state: st.session_state.jd_text     = ""

# ── Sidebar — upload inputs ───────────────────────────────────────────────────
with st.sidebar:
    st.title("📋 AI Resume Assistant")
    st.caption("Score · Rewrite · Cover Letter · Interview Prep")

    st.divider()
    st.markdown("**Step 1 — Upload your resume**")

    upload_tab, paste_tab = st.tabs(["Upload PDF", "Paste text"])
    with upload_tab:
        pdf_file = st.file_uploader("Upload resume PDF", type=["pdf"])
        if pdf_file:
            resume_bytes = pdf_file.read()
            st.session_state.resume_text = extract_pdf_text(resume_bytes)
            st.success(f"Extracted {len(st.session_state.resume_text.split())} words")

    with paste_tab:
        pasted = st.text_area("Paste resume text:", height=150)
        if st.button("Use pasted resume"):
            st.session_state.resume_text = pasted
            st.success("Resume loaded")

    st.divider()
    st.markdown("**Step 2 — Paste job description**")
    jd = st.text_area(
        "Job description:",
        height      = 200,
        placeholder = "Paste the full job description here..."
    )
    if jd:
        st.session_state.jd_text = jd

    st.divider()

    if st.session_state.resume_text:
        st.success(f"Resume: {len(st.session_state.resume_text.split())} words")
    else:
        st.warning("No resume loaded")

    if st.session_state.jd_text:
        st.success(f"JD: {len(st.session_state.jd_text.split())} words")
    else:
        st.info("No job description")

# ── Main ──────────────────────────────────────────────────────────────────────
st.title("📋 AI Resume & Job Application Assistant")
st.caption("Upload resume + job description → get fit score, tailored resume, cover letter, interview prep")

if not st.session_state.resume_text:
    st.info("Upload your resume in the sidebar to get started.")
    st.markdown("""
**What this assistant does:**
- 🎯 Scores how well your resume fits a job (0-100)
- ✍️ Rewrites your resume tailored to the specific role
- 📝 Generates a personalised cover letter
- 🎤 Predicts interview questions with answer frameworks
- 💡 Suggests improvements to your resume
    """)
    st.stop()

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "🎯 Fit Score",
    "✍️ Rewrite Resume",
    "📝 Cover Letter",
    "🎤 Interview Prep",
    "💡 Improve Resume"
])

# ── Tab 1: Fit Score ──────────────────────────────────────────────────────────
with tab1:
    st.subheader("How well do you fit this job?")

    if not st.session_state.jd_text:
        st.warning("Paste the job description in the sidebar first.")
    else:
        if st.button("Score my fit", type="primary"):
            with st.spinner("Analysing job fit..."):
                result = score_job_fit(
                    st.session_state.resume_text,
                    st.session_state.jd_text
                )

            # Score display
            score   = result.get("overall_score", 0)
            verdict = result.get("verdict", "")
            color   = (
                "green" if score >= 75 else
                "orange" if score >= 55 else
                "red"
            )

            col1, col2 = st.columns([1, 3])
            with col1:
                st.metric("Fit Score", f"{score}/100")
                st.markdown(f"**{verdict}**")
            with col2:
                st.progress(score / 100)

            st.divider()

            col1, col2 = st.columns(2)
            with col1:
                st.markdown("**✅ Matching skills**")
                for s in result.get("matching_skills", []):
                    st.markdown(f"- {s}")

                st.markdown("**💪 Your strengths**")
                for s in result.get("strengths", []):
                    st.markdown(f"- {s}")

            with col2:
                st.markdown("**❌ Missing skills**")
                for s in result.get("missing_skills", []):
                    st.markdown(f"- {s}")

                st.markdown("**⚠️ Gaps to address**")
                for g in result.get("gaps", []):
                    st.markdown(f"- {g}")

            st.divider()
            st.markdown("**📌 Recommendations**")
            for r in result.get("recommendations", []):
                st.markdown(f"- {r}")

# ── Tab 2: Rewrite Resume ─────────────────────────────────────────────────────
with tab2:
    st.subheader("Tailored resume for this specific job")

    if not st.session_state.jd_text:
        st.warning("Paste the job description in the sidebar first.")
    else:
        if st.button("Rewrite my resume", type="primary"):
            with st.spinner("Rewriting resume for this role..."):
                rewritten = rewrite_resume(
                    st.session_state.resume_text,
                    st.session_state.jd_text
                )

            st.markdown(rewritten)
            st.divider()

            # Download
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    "Download tailored resume (.md)",
                    rewritten,
                    file_name           = "tailored_resume.md",
                    mime                = "text/markdown",
                    use_container_width = True
                )
            with col2:
                # Build DOCX
                try:
                    from docx import Document
                    doc = Document()
                    doc.add_heading("Tailored Resume", 0)
                    for line in rewritten.split("\n"):
                        if line.startswith("## "):
                            doc.add_heading(line[3:], 2)
                        elif line.startswith("- "):
                            doc.add_paragraph(line[2:], style="List Bullet")
                        elif line.strip():
                            doc.add_paragraph(line)

                    docx_path = "outputs/tailored_resume.docx"
                    doc.save(docx_path)

                    with open(docx_path, "rb") as f:
                        st.download_button(
                            "Download tailored resume (.docx)",
                            f.read(),
                            file_name           = "tailored_resume.docx",
                            mime                = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width = True
                        )
                except Exception:
                    st.info("Install python-docx for Word download")

# ── Tab 3: Cover Letter ───────────────────────────────────────────────────────
with tab3:
    st.subheader("Personalised cover letter")

    if not st.session_state.jd_text:
        st.warning("Paste the job description in the sidebar first.")
    else:
        col1, col2, col3 = st.columns(3)
        with col1:
            company_name = st.text_input(
                "Company name:",
                placeholder = "e.g. TechCorp India"
            )
        with col2:
            hiring_manager = st.text_input(
                "Hiring manager (optional):",
                placeholder = "e.g. Mr. Amith Kumar"
            )
        with col3:
            tone = st.selectbox(
                "Tone:",
                ["professional", "enthusiastic", "concise"]
            )

        if st.button("Generate cover letter", type="primary"):
            with st.spinner("Writing your cover letter..."):
                letter = generate_cover_letter(
                    st.session_state.resume_text,
                    st.session_state.jd_text,
                    company_name,
                    hiring_manager,
                    tone
                )

            st.text_area("Cover letter:", letter, height=400)
            st.download_button(
                "Download cover letter (.txt)",
                letter,
                file_name = "cover_letter.txt",
                mime      = "text/plain"
            )

# ── Tab 4: Interview Prep ─────────────────────────────────────────────────────
with tab4:
    st.subheader("Predicted interview questions + answer frameworks")

    if not st.session_state.jd_text:
        st.warning("Paste the job description in the sidebar first.")
    else:
        if st.button("Predict interview questions", type="primary"):
            with st.spinner("Predicting questions..."):
                questions = predict_interview_questions(
                    st.session_state.resume_text,
                    st.session_state.jd_text
                )

            type_colors = {
                "Technical":     "🔵",
                "Behavioural":   "🟢",
                "Situational":   "🟡",
                "Role-specific": "🔴",
                "Motivational":  "🟣"
            }

            for i, q in enumerate(questions, 1):
                qtype = q.get("type", "General")
                icon  = type_colors.get(qtype, "⚪")

                with st.expander(
                    f"{icon} Q{i}: {q.get('question', '')[:80]}",
                    expanded=(i <= 2)
                ):
                    st.markdown(f"**Type:** {qtype}")
                    st.markdown(f"**Question:** {q.get('question', '')}")
                    st.divider()
                    st.markdown("**Suggested answer framework:**")
                    st.markdown(q.get("answer_framework", ""))

# ── Tab 5: Improve Resume ─────────────────────────────────────────────────────
with tab5:
    st.subheader("Standalone resume improvement analysis")
    st.caption("No job description needed — general resume quality review")

    if st.button("Analyse my resume", type="primary"):
        with st.spinner("Reviewing your resume..."):
            result = improve_resume(st.session_state.resume_text)

        # Score
        score = result.get("overall_score", 0)
        col1, col2 = st.columns([1, 3])
        with col1:
            st.metric("Resume Score", f"{score}/100")
        with col2:
            st.progress(score / 100)
            st.caption(result.get("summary", ""))

        st.divider()

        col1, col2 = st.columns(2)

        with col1:
            st.markdown("**🔄 Weak verbs to replace**")
            for wv in result.get("weak_verbs", []):
                st.markdown(
                    f"- ~~{wv.get('original', '')}~~ → **{wv.get('stronger', '')}**"
                )

            st.markdown("**📊 Add metrics to these**")
            for m in result.get("missing_metrics", []):
                st.markdown(f"- {m}")

        with col2:
            st.markdown("**✅ Top 5 improvements**")
            for i, imp in enumerate(result.get("top_5_improvements", []), 1):
                st.markdown(f"{i}. {imp}")

            st.markdown("**⚙️ ATS issues**")
            for a in result.get("ats_issues", []):
                st.markdown(f"- {a}")

        if result.get("missing_sections"):
            st.markdown("**📌 Consider adding these sections**")
            for s in result["missing_sections"]:
                st.markdown(f"- {s}")